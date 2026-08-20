import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_section(doc, title, content):
    """Add a section heading and content to document"""
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 70, 127)
    doc.add_paragraph(content)
    doc.add_paragraph()

def generate_word_document(bsp_data: dict):
    """
    Takes AI generated content and creates Word BSP document
    """
    
    client = bsp_data["client_details"]
    content = bsp_data["content"]
    plan_type = bsp_data["plan_type"]
    
    # Create document
    doc = Document()
    
    # ===== TITLE =====
    title = doc.add_heading(
        f'{plan_type.upper()} BEHAVIOUR SUPPORT PLAN', 
        level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ===== CLIENT DETAILS TABLE =====
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # Table header
    header_row = table.rows[0]
    header_row.cells[0].text = "Client Information"
    header_row.cells[1].text = "Details"
    
    # Make header bold
    for cell in header_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
    
    # Fill in client details
    details = [
        ("Client Name", client["name"]),
        ("Primary Diagnosis", client["diagnosis"]),
        ("Support Worker", client.get("support_worker", "Not specified")),
        ("BSP Coordinator", client.get("coordinator", "Not specified")),
        ("Date Created", datetime.now().strftime("%d %B %Y")),
    ]
    
    for i, (label, value) in enumerate(details, 1):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # ===== REVIEW DATE =====
    review = doc.add_paragraph()
    review_run = review.add_run(
        f"Review Date: {'3 months' if plan_type == 'interim' else '12 months'} from date of plan"
    )
    review_run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph("_" * 60)
    doc.add_paragraph()
    
    # ===== BSP SECTIONS =====
    for i, (section_name, section_content) in enumerate(content.items(), 1):
        add_section(
            doc, 
            f"{i}. {section_name.upper()}", 
            section_content
        )
    
    # ===== SIGN OFF =====
    doc.add_heading("SIGN OFF", level=1)
    doc.add_paragraph()
    
    signoff_table = doc.add_table(rows=4, cols=2)
    signoff_table.style = 'Table Grid'
    
    signoffs = [
        ("Behaviour Support Practitioner", "________________________"),
        ("Participant / Guardian", "________________________"),
        ("Support Worker", "________________________"),
        ("Date", "________________________"),
    ]
    
    for i, (label, value) in enumerate(signoffs):
        row = signoff_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # ===== DISCLAIMER =====
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "This BSP has been generated based on internal resource documents only. "
        "All content must be reviewed and approved by a qualified Behaviour Support "
        "Practitioner before implementation."
    )
    disclaimer_run.font.italic = True
    disclaimer_run.font.size = Pt(9)
    
    # ===== SAVE DOCUMENT =====
    os.makedirs("output", exist_ok=True)
    
    client_name = client["name"].replace(" ", "_")
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"BSP_{plan_type}_{client_name}_{date_str}.docx"
    output_path = f"output/{filename}"
    
    doc.save(output_path)
    print(f"✓ Document saved: {output_path}")
    
    return output_path, filename